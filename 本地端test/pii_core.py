import os
import re
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

SKIP_DIRS = {
    "#recycle",
    "@eaDir",
    "$RECYCLE.BIN",
    "System Volume Information",
    "_restricted"
}

SUPPORTED_EXTENSIONS = {".txt", ".csv", ".log", ".docx", ".xlsx", ".pdf"}


def detect_pii_types(text):
    pii_types = []

    patterns = {
        "身分證字號": r'[A-Z][12]\d{8}',
        "手機號碼": r'09\d{8}',
        "Email": r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+',
    }

    for pii_name, pattern in patterns.items():
        if re.search(pattern, text):
            pii_types.append(pii_name)

    return pii_types


def read_text_file(file_path):
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def read_docx_file(file_path):
    doc = Document(file_path)
    parts = [p.text for p in doc.paragraphs if p.text]

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                parts.append(" ".join(row_text))

    return "\n".join(parts)


def read_xlsx_file(file_path):
    wb = load_workbook(file_path, read_only=True, data_only=True)
    all_text = []

    for sheet in wb.worksheets:
        all_text.append(f"[工作表]{sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            row_text = [str(cell) for cell in row if cell is not None]
            if row_text:
                all_text.append(" ".join(row_text))

    return "\n".join(all_text)


def read_pdf_file(file_path):
    reader = PdfReader(file_path)
    all_text = []

    for page in reader.pages:
        text = page.extract_text()
        if text:
            all_text.append(text)

    return "\n".join(all_text)


def extract_text(file_path):
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        return None

    if ext in {".txt", ".csv", ".log"}:
        return read_text_file(file_path)
    elif ext == ".docx":
        return read_docx_file(file_path)
    elif ext == ".xlsx":
        return read_xlsx_file(file_path)
    elif ext == ".pdf":
        return read_pdf_file(file_path)

    return None


def scan_file(file_path):
    try:
        text = extract_text(file_path)
        if not text:
            return []

        return detect_pii_types(text)

    except Exception as e:
        return {"error": f"讀檔失敗: {file_path} | {e}"}


def run_scan(config, progress_callback):
    scan_root = config["scan_root"]
    restricted_root = config["restricted_root"]

    if not os.path.exists(scan_root):
        raise Exception(f"掃描路徑不存在：{scan_root}")

    results = []

    for root, dirs, files in os.walk(scan_root, topdown=True):
        dirs[:] = [d for d in dirs if d not in SKIP_DIRS]

        for file_name in files:
            file_path = os.path.join(root, file_name)
            _, ext = os.path.splitext(file_path)

            if ext.lower() not in SUPPORTED_EXTENSIONS:
                continue

            # 避免掃到 restricted 裡的檔案
            if restricted_root and os.path.abspath(file_path).startswith(os.path.abspath(restricted_root)):
                continue

            progress_callback(f"掃描中: {file_path}")

            pii_types = scan_file(file_path)

            if isinstance(pii_types, dict) and "error" in pii_types:
                progress_callback(pii_types["error"])
                continue

            if pii_types:
                try:
                    size_kb = round(os.path.getsize(file_path) / 1024, 2)
                except Exception:
                    size_kb = "N/A"

                results.append({
                    "name": os.path.basename(file_path),
                    "ext": ext.lower(),
                    "path": file_path,
                    "size_kb": size_kb,
                    "pii_types": pii_types,
                    "action_result": "偵測到個資"
                })

    return results